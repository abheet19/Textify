import os
import logging
import requests
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
from wordcloud import WordCloud
from bs4 import BeautifulSoup
import re
from transformers import pipeline, T5Tokenizer
from fpdf import FPDF
from docx import Document
from config import WORDCLOUD_PATH, MODEL_NAME
import hashlib
import datetime

# Set up logging
logging.basicConfig(level=logging.INFO)

# Initialize the summarization model globally to avoid repeated loading
summarizer = pipeline("summarization", model=MODEL_NAME, device=0 if os.getenv('CUDA_AVAILABLE') else -1)
tokenizer = T5Tokenizer.from_pretrained(MODEL_NAME)

def clean_text_and_generate_wordcloud(file_content):
    """
    Clean text and generate word cloud, preserving structure for summarization.
    
    Args:
        file_content (str): The raw input text.

    Returns:
        str: The cleaned text with preserved structure.
    """
    # Create word cloud version (aggressive cleaning)
    wordcloud_text = file_content
    wordcloud_text = re.sub(r'\d', ' ', wordcloud_text)
    wordcloud_text = re.sub(r'\W', ' ', wordcloud_text)
    wordcloud_text = re.sub(r'\s+', ' ', wordcloud_text)
    wordcloud_text = re.sub(r'\[[0-9]*\]', ' ', wordcloud_text)

    # Clean version for actual text processing (preserve structure)
    cleaned_text = file_content
    cleaned_text = re.sub(r'\[[0-9]*\]', ' ', cleaned_text)
    
    # Preserve paragraph structure
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    
    # Clean up spaces within lines, but preserve line/paragraph structure
    lines = cleaned_text.split('\n')
    cleaned_lines = []
    for line in lines:
        clean_line = re.sub(r'\s+', ' ', line).strip()
        cleaned_lines.append(clean_line)
    
    cleaned_text = '\n'.join(cleaned_lines)
    cleaned_text = cleaned_text.strip()

    try:
        os.makedirs(os.path.dirname(WORDCLOUD_PATH), exist_ok=True)
        plt.close('all')
        
        wordcloud = WordCloud(max_font_size=100, max_words=100, background_color="white").generate(wordcloud_text)
        plt.figure()
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.savefig(WORDCLOUD_PATH)
        plt.close()
    except Exception as e:
        logging.error(f"Error generating word cloud: {e}")
    
    return cleaned_text

def generate_bert_summary(cleaned_text):
    """
    Generate an enhanced summary using T5-large model with improved handling for long content.

    Args:
        cleaned_text (str): The cleaned input text.

    Returns:
        str: The generated summary with improved formatting.
    """
    try:
        if not cleaned_text or len(cleaned_text.strip()) == 0:
            return "No content to summarize."
        
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text.strip())
        
        if len(cleaned_text.split()) < 30:
            return "Text too short for summarization. Please provide more content."
        
        words = cleaned_text.split()
        
        if len(words) > 800:
            # Count articles for adaptive summarization
            news_story_count = cleaned_text.count('NEWS STORY:')
            headline_count = cleaned_text.count('HEADLINE:')
            article_markers = news_story_count + headline_count
            
            paragraphs = [p.strip() for p in cleaned_text.split('\n\n') if len(p.strip()) > 30]
            lines_with_content = [line.strip() for line in cleaned_text.split('\n') if len(line.strip()) > 50]
            
            # Enhanced article detection for India Today
            if 'indiatoday' in cleaned_text.lower() or len(lines_with_content) > 10:
                article_count = max(
                    article_markers,
                    len(paragraphs),
                    len(lines_with_content) // 2,
                    15
                )
            else:
                article_count = max(article_markers, len(paragraphs) // 3, 1)
            
            # Adaptive parameters based on article count
            if article_count >= 8:
                max_tokens_per_chunk = 150
                min_length_per_chunk = 60
                final_max_tokens = 500
                final_min_length = 300
                summary_style = "Create a detailed news summary organized by topics with clear headings. Group related stories under categories like 'Asia Cup:', 'International News:', 'Domestic News:', 'Business:', etc. Use bullet points and subheadings for clarity"
            elif article_count >= 4:
                max_tokens_per_chunk = 120
                min_length_per_chunk = 50
                final_max_tokens = 350
                final_min_length = 200
                summary_style = "Create a structured news summary with topic headings. Group similar stories together under clear categories"
            else:
                max_tokens_per_chunk = 80
                min_length_per_chunk = 30
                final_max_tokens = 200
                final_min_length = 80
                summary_style = "Create a concise summary covering the key news topics mentioned"
            
            # Smart chunking strategy
            chunks = []
            chunk_size = 400
            overlap = 50
            
            for i in range(0, len(words), chunk_size - overlap):
                chunk = ' '.join(words[i:i + chunk_size])
                chunks.append(chunk)
                if i + chunk_size >= len(words):
                    break
            
            # Summarize each chunk
            chunk_summaries = []
            for i, chunk in enumerate(chunks[:3]):
                try:
                    input_ids = tokenizer.encode(f"Provide a detailed summary covering all key points and topics mentioned: {chunk}", truncation=True, max_length=512)
                    chunk_text = tokenizer.decode(input_ids, skip_special_tokens=True)
                    
                    result = summarizer(chunk_text, max_new_tokens=max_tokens_per_chunk, min_length=min_length_per_chunk, do_sample=False)
                    chunk_summary = result[0]["summary_text"]
                    chunk_summaries.append(chunk_summary)
                except Exception as e:
                    continue
            
            if chunk_summaries:
                combined_text = '\n\n'.join(chunk_summaries)
                
                if len(combined_text.split()) > 80:
                    try:
                        final_input = f"{summary_style}: {combined_text}"
                        result = summarizer(final_input, max_new_tokens=final_max_tokens, min_length=final_min_length, do_sample=False)
                        summary = result[0]["summary_text"]
                        
                        # Clean up prompt leakage
                        if summary.startswith(summary_style.split(':')[0]):
                            summary = summary[len(summary_style.split(':')[0]):].strip()
                        if summary.startswith(":"):
                            summary = summary[1:].strip()
                        
                        # Remove garbled text
                        summary = re.sub(r'[­\-–—]{2,}', '', summary)
                        summary = re.sub(r'\.{3,}', '.', summary)
                        summary = re.sub(r'\s[a-z]\s', ' ', summary)
                        summary = re.sub(r'^[^A-Z]*', '', summary)
                        summary = re.sub(r'\s+', ' ', summary)
                        
                        # Enhanced structure formatting for ALL summary sizes
                        if article_count >= 8:
                            # Extract meaningful sentences
                            sentences = []
                            for s in summary.split('.'):
                                s = s.strip()
                                if len(s) > 15 and any(c.isalpha() for c in s):
                                    if not s.endswith('.'):
                                        s += '.'
                                    sentences.append(s)
                            
                            if len(sentences) >= 3:
                                structured_summary = ""
                                
                                # Categorize sentences by keywords
                                sports_keywords = ['cricket', 'cup', 'match', 'player', 'team', 'sport', 'asia cup', 'pakistan cricket', 'batting', 'bowling', 'wicket', 'runs', 'tournament']
                                international_keywords = ['russia', 'nato', 'foreign ministry', 'pressure', 'oil', 'moscow', 'us', 'united states', 'ties with india', 'international relations']
                                crime_keywords = ['killed', 'murder', 'lover', 'husband', 'dumping body', 'crime', 'police', 'arrested', 'victim', 'criminal']
                                domestic_keywords = ['nepal', 'president', 'minister', 'government', 'cabinet', 'domestic policy', 'national']
                                
                                sports_sentences = []
                                international_sentences = []
                                crime_sentences = []
                                domestic_sentences = []
                                other_sentences = []
                                
                                for sentence in sentences:
                                    sentence_lower = sentence.lower()
                                    sentence = re.sub(r'[­\-–—]{2,}', '', sentence)
                                    sentence = re.sub(r'\s+', ' ', sentence).strip()
                                    
                                    if len(sentence) > 15:
                                        if any(keyword in sentence_lower for keyword in crime_keywords):
                                            crime_sentences.append(sentence)
                                        elif any(keyword in sentence_lower for keyword in sports_keywords) and ('cricket' in sentence_lower or 'cup' in sentence_lower or 'match' in sentence_lower):
                                            sports_sentences.append(sentence)
                                        elif any(keyword in sentence_lower for keyword in international_keywords):
                                            international_sentences.append(sentence)
                                        elif any(keyword in sentence_lower for keyword in domestic_keywords):
                                            domestic_sentences.append(sentence)
                                        else:
                                            other_sentences.append(sentence)
                                
                                # Build structured output with HTML formatting
                                if sports_sentences:
                                    structured_summary += "<strong>Sports & Asia Cup:</strong><br>"
                                    for sentence in sports_sentences[:3]:
                                        structured_summary += f"• {sentence}<br>"
                                    structured_summary += "<br>"
                                
                                if international_sentences:
                                    structured_summary += "<strong>International News:</strong><br>"
                                    for sentence in international_sentences[:3]:
                                        structured_summary += f"• {sentence}<br>"
                                    structured_summary += "<br>"
                                
                                if crime_sentences:
                                    structured_summary += "<strong>Crime & Security:</strong><br>"
                                    for sentence in crime_sentences[:3]:
                                        structured_summary += f"• {sentence}<br>"
                                    structured_summary += "<br>"
                                
                                if domestic_sentences:
                                    structured_summary += "<strong>Domestic News:</strong><br>"
                                    for sentence in domestic_sentences[:3]:
                                        structured_summary += f"• {sentence}<br>"
                                    structured_summary += "<br>"
                                
                                if other_sentences:
                                    structured_summary += "<strong>Other News:</strong><br>"
                                    for sentence in other_sentences[:2]:
                                        structured_summary += f"• {sentence}<br>"
                                
                                summary = structured_summary.strip()
                            else:
                                summary = re.sub(r'[­\-–—\.]{2,}', '', summary)
                                summary = re.sub(r'\s+', ' ', summary).strip()
                        
                        elif article_count >= 4:
                            # Create paragraph breaks for medium summaries
                            if '. ' in summary:
                                sentences = [s.strip() for s in summary.split('. ') if s.strip() and len(s.strip()) > 10]
                                if len(sentences) >= 2:
                                    paragraphs = []
                                    sentences_per_para = max(1, len(sentences) // 3)
                                    
                                    for i in range(0, len(sentences), sentences_per_para):
                                        para_sentences = sentences[i:i + sentences_per_para]
                                        para = '. '.join(para_sentences)
                                        if not para.endswith('.'):
                                            para += '.'
                                        paragraphs.append(para)
                                    
                                    summary = '<br><br>'.join(paragraphs)
                                    
                        else:
                            # Ensure at least 2 paragraphs for small summaries
                            if '. ' in summary:
                                sentences = [s.strip() for s in summary.split('. ') if s.strip() and len(s.strip()) > 10]
                                if len(sentences) >= 2:
                                    mid_point = len(sentences) // 2
                                    para1 = '. '.join(sentences[:mid_point])
                                    para2 = '. '.join(sentences[mid_point:])
                                    
                                    if not para1.endswith('.'):
                                        para1 += '.'
                                    if not para2.endswith('.'):
                                        para2 += '.'
                                    
                                    summary = f"{para1}<br><br>{para2}"
                                    
                            summary = re.sub(r'[­\-–—\.]{2,}', '', summary)
                            summary = re.sub(r'\s+', ' ', summary).strip()
                        
                    except:
                        summary = combined_text.replace('. ', '.\n\n').strip()
                else:
                    summary = combined_text
            else:
                # Fallback for no chunk summaries
                input_ids = tokenizer.encode(cleaned_text, truncation=True, max_length=800)
                fallback_text = tokenizer.decode(input_ids, skip_special_tokens=True)
                input_text = f"{summary_style}: {fallback_text}"
                result = summarizer(input_text, max_new_tokens=final_max_tokens, min_length=final_min_length, do_sample=False)
                summary = result[0]["summary_text"]
                summary = summary.replace('. ', '.\n\n').strip()
        else:
            # For shorter content, use direct summarization
            input_ids = tokenizer.encode(f"summarize: {cleaned_text}", truncation=True, max_length=512)
            input_text = tokenizer.decode(input_ids, skip_special_tokens=True)
            
            result = summarizer(input_text, max_new_tokens=120, min_length=40, do_sample=False)
            summary = result[0]["summary_text"]
        
        # Final cleanup
        if '<br>' in summary:
            lines = summary.split('<br>')
        else:
            lines = summary.split('\n')
        
        clean_lines = []
        for line in lines:
            line = line.strip()
            line = re.sub(r'enen\s*-en\s*en-ena\s*aen.*', '', line)
            line = re.sub(r'it says\s*[­\-–—\.a\s]*', '', line)
            line = re.sub(r'[­\-–—]{2,}', '', line)
            line = re.sub(r'\.{3,}', '.', line)
            line = re.sub(r'\s+', ' ', line).strip()
            
            if (len(line) > 3 and 
                not re.match(r'^[^a-zA-Z]*$', line) and
                not re.search(r'^[a-z\s\.\-–—]*$', line) and
                'enen' not in line and
                len([c for c in line if c.isalpha()]) > 3):
                clean_lines.append(line)
        
        if '<br>' in summary:
            summary = '<br>'.join(clean_lines)
        else:
            summary = '\n'.join(clean_lines)
        
        summary = summary.strip()
        
        return summary
        
    except Exception as e:
        logging.error(f"Error generating summary: {e}")
        return "Error generating summary. Please try again."

def fetch_article(url):
    """
    Fetch and extract text content from a URL with enhanced scraping for news homepages.
    
    Args:
        url (str): The URL to fetch content from.
        
    Returns:
        str: The extracted text content.
    """
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove unwanted elements
        for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 'advertisement']):
            element.decompose()
        
        # Check if this is a homepage or news aggregation page
        is_homepage = (
            url.endswith('.com') or url.endswith('.com/') or 
            url.endswith('.in') or url.endswith('.in/') or
            'indiatoday.in' in url or
            len(soup.find_all('h1')) > 3 or
            len(soup.find_all('h2')) > 5
        )
        
        if is_homepage:
            # Enhanced strategy for homepages/news aggregation
            news_content = []
            
            # 1. Extract headlines
            for i, headline in enumerate(soup.find_all(['h1', 'h2', 'h3'])[:15]):
                headline_text = headline.get_text(strip=True)
                if headline_text and 20 < len(headline_text) < 200:
                    news_content.append(f"HEADLINE: {headline_text}")
            
            # 2. Extract article blocks
            articles = soup.find_all('article')
            for i, article in enumerate(articles[:12]):
                article_text = article.get_text(separator=' ', strip=True)
                if article_text and 30 < len(article_text.split()) < 150:
                    news_content.append(f"ARTICLE: {article_text}")
            
            # 3. Extract news paragraphs
            news_paragraphs = []
            for p in soup.find_all('p'):
                p_text = p.get_text(strip=True)
                if p_text and 15 < len(p_text.split()) < 100:
                    if any(keyword in p_text.lower() for keyword in ['said', 'according', 'reported', 'news', 'today', 'announced']):
                        news_paragraphs.append(p_text)
            
            for para in news_paragraphs[:8]:
                news_content.append(f"NEWS STORY: {para}")
            
            # 4. Extract story links text
            for link in soup.find_all('a', href=True)[:20]:
                link_text = link.get_text(strip=True)
                if link_text and 10 < len(link_text.split()) < 50:
                    if not any(skip in link_text.lower() for skip in ['click', 'read more', 'subscribe', 'login', 'register']):
                        news_content.append(f"STORY LINK: {link_text}")
            
            # Combine all content
            text = '\n\n'.join(news_content)
            
        else:
            # Original strategy for single articles
            articles = soup.find_all("article")
            article_texts = [a.get_text(separator=" ", strip=True) for a in articles if a.get_text(strip=True)]
            
            headlines = []
            for tag in soup.find_all(['h1', 'h2']):
                headline = tag.get_text(strip=True)
                if headline:
                    headlines.append(headline)
            
            paragraphs = [p.get_text(strip=True) for p in soup.find_all("p") if p.get_text(strip=True)]
            
            # Combine content
            all_content = []
            all_content.extend(headlines)
            all_content.extend(article_texts)
            all_content.extend(paragraphs)
            
            text = ' '.join(all_content)
        
        if not text or len(text.strip()) < 100:
            # Fallback: get all visible text
            text = soup.get_text(separator=' ', strip=True)
        
        return text
        
    except Exception as e:
        logging.error(f"Error fetching article: {e}")
        return f"Error fetching content: {e}"

def summarize_url(url):
    """
    Main function to summarize content from a URL.
    
    Args:
        url (str): The URL to summarize.
        
    Returns:
        str: The summary of the URL content.
    """
    try:
        article_text = fetch_article(url)
        
        if article_text.startswith("Error"):
            return article_text
        
        cleaned_text = clean_text_and_generate_wordcloud(article_text)
        
        summary_text = generate_bert_summary(cleaned_text)
        
        if summary_text.startswith("Error") or summary_text.startswith("Text too short"):
            return summary_text
        
        return summary_text
    except Exception as e:
        logging.error(f"Error summarizing URL: {e}")
        return f"Error processing URL: {e}"

def generate_pdf_report(summary_text, filename=None):
    """
    Generate a PDF report of the summary.
    
    Args:
        summary_text (str): The summary text to include in the PDF.
        filename (str, optional): Custom filename for the PDF.
        
    Returns:
        str: The path to the generated PDF file.
    """
    if not filename:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Text_Summary_Report_{timestamp}.pdf"
    
    pdf_path = os.path.join("downloads", filename)
    os.makedirs("downloads", exist_ok=True)
    
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt="Text Summarization Report", ln=True, align='C')
        pdf.ln(10)
        
        # Add summary content
        pdf.set_font("Arial", size=12)
        
        # Handle HTML formatting in summary
        clean_summary = summary_text.replace('<strong>', '').replace('</strong>', '').replace('<br>', '\n')
        
        # Split text into lines to fit PDF width
        lines = clean_summary.split('\n')
        for line in lines:
            if len(line) > 80:
                words = line.split(' ')
                current_line = ""
                for word in words:
                    if len(current_line + word) < 80:
                        current_line += word + " "
                    else:
                        pdf.cell(200, 10, txt=current_line.strip(), ln=True)
                        current_line = word + " "
                if current_line:
                    pdf.cell(200, 10, txt=current_line.strip(), ln=True)
            else:
                pdf.cell(200, 10, txt=line, ln=True)
        
        pdf.output(pdf_path)
        return pdf_path
        
    except Exception as e:
        logging.error(f"Error generating PDF: {e}")
        return None

def generate_docx_report(summary_text, filename=None):
    """
    Generate a DOCX report of the summary.
    
    Args:
        summary_text (str): The summary text to include in the DOCX.
        filename (str, optional): Custom filename for the DOCX.
        
    Returns:
        str: The path to the generated DOCX file.
    """
    if not filename:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Text_Summary_Report_{timestamp}.docx"
    
    docx_path = os.path.join("downloads", filename)
    os.makedirs("downloads", exist_ok=True)
    
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading('Text Summarization Report', 0)
        
        # Add summary content
        clean_summary = summary_text.replace('<strong>', '').replace('</strong>', '').replace('<br>', '\n')
        
        lines = clean_summary.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())
        
        doc.save(docx_path)
        return docx_path
        
    except Exception as e:
        logging.error(f"Error generating DOCX: {e}")
        return None
