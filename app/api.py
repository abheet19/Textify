"""
API blueprint for the Text Summarizer System.

This module defines optimized routes for handling file uploads, URL submissions,
summary generation, and file downloads.

OPTIMIZATIONS IMPLEMENTED:
- Consolidated error handling with standardized handle_error() function
- Generic process_summary_request() function to reduce code duplication
- Unified download route supporting both DOCX and PDF formats
- Removed redundant imports and improved code organization
- Backward compatibility maintained for existing download routes
"""

import os
import logging
from datetime import datetime
from flask import Blueprint, request, render_template, send_file
from .service import (
    clean_text_and_generate_wordcloud,
    generate_bert_summary,
    summarize_url,
    generate_pdf_report,
    generate_docx_report
)

# File processing imports
import PyPDF2
from docx import Document
import io

# Configure logging for the API module
logging.basicConfig(level=logging.ERROR)

api = Blueprint('api', __name__)

def handle_error(error_message, status_code=500):
    """
    Standardized error handling for API routes.
    
    Args:
        error_message (str): The error message to display
        status_code (int): HTTP status code
        
    Returns:
        tuple: Error response and status code
    """
    logging.error(error_message)
    
    # Create a user-friendly error page
    error_html = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Error - Text Summarizer</title>
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
        <style>
            body {{ background-color: #f8f9fa; }}
            .error-container {{ 
                min-height: 100vh; 
                display: flex; 
                align-items: center; 
                justify-content: center; 
            }}
            .error-card {{ 
                max-width: 600px; 
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1); 
            }}
            .error-icon {{ 
                font-size: 4rem; 
                color: #dc3545; 
            }}
        </style>
    </head>
    <body>
        <div class="error-container">
            <div class="error-card card">
                <div class="card-body text-center p-5">
                    <div class="error-icon mb-4">⚠️</div>
                    <h2 class="card-title text-danger mb-3">File Processing Error</h2>
                    <p class="card-text lead mb-4">{error_message}</p>
                    <div class="mb-4">
                        <h5>What you can try:</h5>
                        <ul class="text-left">
                            <li>Make sure your file is not corrupted or empty</li>
                            <li>For PDF files: Ensure they contain readable text (not just images)</li>
                            <li>For Word files: Save as .docx format (not .doc)</li>
                            <li>For text files: Ensure they are properly encoded (UTF-8)</li>
                            <li>Try uploading a different file</li>
                        </ul>
                    </div>
                    <div>
                        <a href="/PDF" class="btn btn-primary mr-2">Try Another File</a>
                        <a href="/" class="btn btn-secondary">Go Home</a>
                    </div>
                </div>
            </div>
        </div>
    </body>
    </html>
    """
    
    return error_html, status_code

def process_summary_request(input_data, input_type="file"):
    """
    Generic summary processing function to reduce code duplication.
    
    Args:
        input_data: File object or URL string
        input_type (str): Type of input - "file" or "url"
        
    Returns:
        tuple: (success, summary_text_or_error_message)
    """
    try:
        if input_type == "file":
            if not input_data:
                return False, "Please provide a file - No file was uploaded"
            
            if not input_data.filename or input_data.filename == '':
                return False, "Please select a file with a valid filename - The uploaded file has no name"
                
            filename = input_data.filename.strip()
            logging.info(f"Processing file upload: {filename}")
            
            # Process file directly from memory (no need to save to disk)
            try:
                file_content = process_file(input_data)
                logging.info(f"File processed successfully, content length: {len(file_content)}")
                
                cleaned_text = clean_text_and_generate_wordcloud(file_content)
                summary_text = generate_bert_summary(cleaned_text)
                
                logging.info("Summary generated successfully")
                return True, summary_text
                
            except ValueError as e:
                # These are user-friendly errors from process_file
                logging.error(f"File processing validation error: {e}")
                return False, str(e)
            except Exception as e:
                # Unexpected errors
                logging.error(f"Unexpected file processing error: {e}")
                return False, f"Unable to process file '{filename}': {str(e)}"
                
        elif input_type == "url":
            if not input_data or not input_data.strip():
                return False, "Please provide a valid URL"
                
            logging.info(f"Processing URL: {input_data}")
            summary_text = summarize_url(input_data)
            
            if summary_text == "Invalid URL":
                return False, "Error! Please provide a correct URL - The URL could not be accessed or processed"
            
            logging.info("URL summary generated successfully")
            return True, summary_text
            
    except Exception as e:
        logging.error(f"Unexpected error in process_summary_request: {e}")
        return False, f"An unexpected error occurred while processing your {input_type}: {str(e)}"
    
    return False, f"Unknown {input_type} processing error - Please try again"

def process_file(file_obj):
    """
    Process uploaded file and extract text content.
    
    Args:
        file_obj: Flask file object
        
    Returns:
        str: Extracted text content
    """
    try:
        # Basic validation
        if not file_obj:
            raise ValueError("No file object provided")
            
        if not hasattr(file_obj, 'filename'):
            raise ValueError("Invalid file object - missing filename attribute")
            
        if not file_obj.filename:
            raise ValueError("No filename provided - please select a file")
        
        filename = file_obj.filename.strip()
        if not filename:
            raise ValueError("Empty filename - please select a valid file")
            
        filename_lower = filename.lower()
        logging.info(f"Processing file: {filename} (type check: {filename_lower})")
        
        # Read file data once and validate size
        file_data = file_obj.read()
        if not file_data:
            raise ValueError(f"File '{filename}' is empty - please upload a file with content")
        
        file_size = len(file_data)
        logging.info(f"File size: {file_size} bytes")
        
        # Check minimum file size (avoid processing tiny files)
        if file_size < 10:
            raise ValueError(f"File '{filename}' is too small ({file_size} bytes) - please upload a valid file")
        
        # Check file extension and process accordingly
        if filename_lower.endswith('.txt'):
            logging.info("Processing as TXT file")
            try:
                if isinstance(file_data, bytes):
                    # Try different encodings
                    for encoding in ['utf-8', 'utf-16', 'latin1', 'cp1252']:
                        try:
                            text = file_data.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        raise ValueError("Unable to decode text file. Please ensure it's properly encoded.")
                else:
                    text = str(file_data)
                
                text = text.strip()
                if not text:
                    raise ValueError("Text file appears to be empty after processing")
                    
                logging.info(f"TXT file processed successfully, length: {len(text)}")
                return text
            except UnicodeDecodeError:
                raise ValueError("Unable to decode text file. Please ensure it's UTF-8 encoded.")
            
        elif filename_lower.endswith('.pdf'):
            logging.info("Processing as PDF file")
            try:
                # Validate PDF header
                if not file_data.startswith(b'%PDF'):
                    raise ValueError("File is not a valid PDF - missing PDF header")
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                
                # Check if PDF has pages
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF file has no pages")
                
                text = ""
                pages_processed = 0
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                            pages_processed += 1
                    except Exception as page_error:
                        logging.warning(f"Could not extract text from page {page_num + 1}: {page_error}")
                        continue
                
                text = text.strip()
                if not text:
                    raise ValueError("PDF file appears to contain no extractable text. It may be image-based or corrupted.")
                
                logging.info(f"PDF processed successfully, {pages_processed} pages, text length: {len(text)}")
                return text
                
            except PyPDF2.errors.PdfReadError as pdf_error:
                logging.error(f"PDF read error: {pdf_error}")
                raise ValueError(f"Cannot read PDF file - it may be corrupted, password-protected, or not a valid PDF")
            except Exception as pdf_error:
                logging.error(f"PDF processing error: {pdf_error}")
                if "empty file" in str(pdf_error).lower():
                    raise ValueError("PDF file is empty or corrupted")
                elif "not a valid PDF" in str(pdf_error):
                    raise ValueError("File is not a valid PDF format")
                else:
                    raise ValueError(f"Error processing PDF file: {str(pdf_error)}")
            
        elif filename_lower.endswith('.docx'):
            logging.info("Processing as DOCX file")
            try:
                # Validate DOCX file (it's actually a ZIP file)
                if not file_data.startswith(b'PK'):
                    raise ValueError("File is not a valid DOCX - missing ZIP header")
                
                # Try to open as DOCX
                doc = Document(io.BytesIO(file_data))
                text = ""
                paragraphs_processed = 0
                
                for paragraph in doc.paragraphs:
                    if paragraph.text and paragraph.text.strip():
                        text += paragraph.text.strip() + "\n"
                        paragraphs_processed += 1
                
                # Also try to extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                text += cell.text.strip() + " "
                
                text = text.strip()
                if not text:
                    raise ValueError("DOCX file appears to be empty or contains no text")
                
                logging.info(f"DOCX processed successfully, {paragraphs_processed} paragraphs, text length: {len(text)}")
                return text
                
            except Exception as docx_error:
                logging.error(f"DOCX processing error: {docx_error}")
                if "not a zip file" in str(docx_error).lower():
                    raise ValueError("File is not a valid DOCX format - DOCX files must be properly formatted Word documents")
                elif "bad zip file" in str(docx_error).lower():
                    raise ValueError("DOCX file is corrupted or incomplete")
                else:
                    raise ValueError(f"Error processing DOCX file: {str(docx_error)}")
            
        elif filename_lower.endswith('.doc'):
            logging.info("Processing as DOC file (limited support)")
            try:
                # DOC files are more complex, try basic text extraction
                text = file_data.decode('utf-8', errors='ignore')
                
                # Clean up the text (DOC files have a lot of binary data)
                import re
                text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)  # Keep only printable ASCII
                text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                text = text.strip()
                
                if not text or len(text) < 50:
                    raise ValueError("DOC file processing failed - please convert to DOCX format for better results")
                
                logging.info(f"DOC processed (basic extraction), text length: {len(text)}")
                return text
                
            except Exception as doc_error:
                logging.error(f"DOC processing error: {doc_error}")
                raise ValueError("Unable to process .doc file. Please convert to .docx or .txt format for better compatibility.")
        else:
            supported_formats = ['.txt', '.pdf', '.docx', '.doc']
            raise ValueError(f"Unsupported file type '{filename}'. Supported formats: {', '.join(supported_formats)}")
            
    except Exception as e:
        error_msg = str(e)
        logging.error(f"File processing failed: {error_msg}")
        
        # Re-raise with clear error message
        if any(phrase in error_msg for phrase in ["Unsupported file type", "No file", "Empty filename", "Invalid file", "not a valid", "corrupted", "empty"]):
            raise ValueError(error_msg)
        else:
            raise ValueError(f"Error processing file: {error_msg}")

def safe_remove_and_render(template_name, render_func, **kwargs):
    """
    Safely render template and remove temporary files.
    
    Args:
        template_name (str): Template name to render
        render_func: Render function (usually render_template)
        **kwargs: Additional arguments for template rendering
        
    Returns:
        Response: Rendered template response
    """
    remove_files()
    return render_func(template_name, **kwargs)

def remove_files():
    """
    Remove temporary files from static directories.
    """
    try:
        # Remove word cloud images
        static_dir = os.path.join(os.path.dirname(__file__), 'static', 'img', 'wordcloud')
        if os.path.exists(static_dir):
            for file in os.listdir(static_dir):
                if file.endswith(('.png', '.jpg', '.jpeg')):
                    os.remove(os.path.join(static_dir, file))
        
        # Remove temporary download files
        download_dir = os.path.join(os.path.dirname(__file__), 'static', 'download')
        if os.path.exists(download_dir):
            for file in os.listdir(download_dir):
                if file.endswith(('.docx', '.pdf')):
                    os.remove(os.path.join(download_dir, file))
    except Exception as e:
        logging.warning(f"Error removing files: {str(e)}")

def generate_docx(summary_text):
    """
    Generate DOCX file from summary text.
    
    Args:
        summary_text (str): Summary text to convert
        
    Returns:
        str: Path to generated DOCX file
    """
    return generate_docx_report(summary_text)

def generate_pdf(summary_text):
    """
    Generate PDF file from summary text.
    
    Args:
        summary_text (str): Summary text to convert
        
    Returns:
        str: Path to generated PDF file
    """
    return generate_pdf_report(summary_text)

@api.route('/', methods=['GET', 'POST'])
def home():
    """
    Render the home page and remove previous word cloud and DOCX files.

    Returns:
        Response: The rendered home page.
    """
    return safe_remove_and_render('index.html', render_template)

@api.route('/PDF', methods=['GET', 'POST'])
def PDF():
    """
    Render the PDF upload page and remove previous word cloud and DOCX files.

    Returns:
        Response: The rendered PDF upload page.
    """
    return safe_remove_and_render('PDF.html', render_template)

@api.route('/debug_upload', methods=['POST'])
def debug_upload():
    """Debug route to test file upload functionality"""
    try:
        logging.info(f"Debug upload - request.files: {request.files}")
        logging.info(f"Debug upload - request.form: {request.form}")
        
        for key, file in request.files.items():
            logging.info(f"File key: {key}, file: {file}, filename: {getattr(file, 'filename', 'NO_FILENAME')}")
        
        uploaded_file = request.files.get('name')
        if uploaded_file:
            logging.info(f"File found: {uploaded_file.filename}")
            
            # Check file info
            filename = uploaded_file.filename
            content_type = getattr(uploaded_file, 'content_type', 'Unknown')
            
            # Try to read file data
            file_data = uploaded_file.read()
            file_size = len(file_data)
            
            # Reset file pointer for processing
            uploaded_file.seek(0)
            
            logging.info(f"File details - Name: {filename}, Size: {file_size} bytes, Content-Type: {content_type}")
            
            # Check file header
            if file_data:
                header = file_data[:20]
                logging.info(f"File header: {header}")
                
                # Check specific file types
                if filename.lower().endswith('.pdf'):
                    is_pdf = file_data.startswith(b'%PDF')
                    logging.info(f"PDF check - starts with %PDF: {is_pdf}")
                elif filename.lower().endswith('.docx'):
                    is_docx = file_data.startswith(b'PK')
                    logging.info(f"DOCX check - starts with PK (ZIP): {is_docx}")
            
            # Try to process the file
            result = process_file(uploaded_file)
            return f"Success! File processed: {len(result)} characters extracted"
        else:
            return "No file found with key 'name'"
            
    except Exception as e:
        logging.error(f"Debug upload error: {e}")
        return f"Error: {str(e)}"

@api.route('/test_file_info', methods=['POST'])
def test_file_info():
    """Test route to get detailed file information"""
    try:
        response = {"files": [], "form_data": dict(request.form)}
        
        for key, file in request.files.items():
            if file:
                file_data = file.read()
                file.seek(0)  # Reset for potential reuse
                
                file_info = {
                    "form_key": key,
                    "filename": file.filename,
                    "content_type": getattr(file, 'content_type', 'Unknown'),
                    "size": len(file_data),
                    "header": file_data[:20].hex() if file_data else "empty",
                    "is_empty": len(file_data) == 0,
                    "has_pdf_header": file_data.startswith(b'%PDF') if file_data else False,
                    "has_zip_header": file_data.startswith(b'PK') if file_data else False
                }
                response["files"].append(file_info)
        
        return response
        
    except Exception as e:
        return {"error": str(e)}

@api.route('/PDF_result', methods=['GET', 'POST'])
def PDF_result():
    """
    Process the uploaded PDF/DOCX/TXT file, generate a summary, and render the result page.

    Returns:
        Response: The rendered result page with the summary or an error message.
    """
    remove_files()
    
    if request.method == 'POST':
        logging.info(f"POST request received, files: {request.files}")
        uploaded_file = request.files.get('name')
        logging.info(f"Uploaded file: {uploaded_file}")
        if uploaded_file:
            logging.info(f"File filename: {uploaded_file.filename}")
            logging.info(f"File content type: {getattr(uploaded_file, 'content_type', 'Unknown')}")
        
        success, result = process_summary_request(uploaded_file, "file")
        
        if not success:
            return handle_error(result, 400)
            
        return render_template('PDF_result.html', summary=result)
    
    return render_template('PDF.html')

@api.route('/RAW', methods=['GET', 'POST'])
def RAW():
    """
    Render the URL input page and remove previous word cloud and DOCX files.

    Returns:
        Response: The rendered URL input page.
    """
    return safe_remove_and_render('RAW.html', render_template)

@api.route('/RAW_result', methods=['GET', 'POST'])
def RAW_result():
    """
    Process the provided URL, generate a summary, and render the result page.

    Returns:
        Response: The rendered result page with the summary or an error message.
    """
    remove_files()
    
    if request.method == 'POST':
        url = request.form.get('name')
        success, result = process_summary_request(url, "url")
        
        if not success:
            return handle_error(result, 400)
            
        return render_template('RAW_result.html', summary=result)
    
    return render_template('RAW.html')

@api.route('/download/<format_type>', methods=['POST'])
def download_file(format_type):
    """
    Universal download route for both DOCX and PDF formats.
    
    Args:
        format_type (str): Either 'docx' or 'pdf'

    Returns:
        Response: The requested file for download or an error message.
    """
    summary = request.form.get('summary')
    if not summary:
        return handle_error("No summary available to download", 400)
    
    # Validate format
    if format_type not in ['docx', 'pdf']:
        return handle_error("Invalid download format", 400)
    
    try:
        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        download_name = f"Text_Summary_Report_{timestamp}.{format_type}"
        
        # Generate file based on format
        if format_type == 'docx':
            file_path = generate_docx(summary)
        else:  # pdf
            file_path = generate_pdf(summary)
            
        return send_file(file_path, as_attachment=True, download_name=download_name)
        
    except Exception as e:
        return handle_error(f"Error generating {format_type.upper()} file: {str(e)}", 500)

# Maintain backward compatibility with existing routes
@api.route('/download', methods=['POST'])
def download():
    """Legacy DOCX download route - redirects to new unified route."""
    return download_file('docx')

@api.route('/download_pdf', methods=['POST'])
def download_pdf():
    """Legacy PDF download route - redirects to new unified route."""
    return download_file('pdf')
